def get_files_in_folder(folder,extension):
	import os
	file_list = []
	for file in os.listdir(folder):
		if file.endswith(extension):
			file_list.append(os.path.join(folder,file))
	return file_list
			# return os.path.join(folder, file)



def readlines(fname):
	with open(fname) as f:
		return f.read().splitlines()

def move_table_after(table, paragraph):
	tbl, p = table._tbl, paragraph._p
	p.addnext(tbl)

def get_table_from_xls_to_word(kodi,folder):
	from win32com import client
	import os
	excel = client.Dispatch("Excel.Application")
	word = client.Dispatch("Word.Application")
	filename_excel = os.path.join(folder,kodi) + "_preventiv.xlsx"
	filename_word = os.path.join(folder,kodi) + ".docx"
	doc = word.Documents.Add()
	# doc = word.Documents.Open(filename_word)
	book = excel.Workbooks.Open(filename_excel)
	sheet = book.Worksheets("RaportWidth")
	sheet.Range("B6:E12").Copy()
	doc.Content.PasteExcelTable(False,False,False)
	doc.SaveAs(filename_word)
	doc.Close()
	book.Close()
	# word.Quit()

def quit_word_app():
	from win32com import client
	word = client.Dispatch("Word.Application")
	word.Quit()

def get_name(Kodi):
	import os
	csv_folder = r"F:\Dropbox\Projekte UTS-01\DLDP - Venddepozitime\Code"
	csv_file = r"Relacion.csv"
	my_data = readlines(os.path.join(csv_folder,csv_file))
	splitted_data = []
	for data in my_data:
		splitted_data.append(data.split(","))
	koment2 = ""
	for data in splitted_data:
		if str(data[0])==Kodi:
			if str(data[9])=="":
				if int(round(float(data[0]))) > 200000:
					return "Venddepozitim nga MZHU"
				else:
					return "Venddepozitim i panjehesuar"
			else:
				return str(data[9])
	return koment2

def get_sip(Kodi):
	import os
	csv_folder = r"F:\Dropbox\Projekte UTS-01\DLDP - Venddepozitime\Code"
	csv_file = r"Relacion.csv"
	my_data = readlines(os.path.join(csv_folder,csv_file))
	splitted_data = []
	for data in my_data:
		splitted_data.append(data.split(","))
	koment2 = ""
	for data in splitted_data:
		if str(data[0])==Kodi:
			koment2 = "Ky venddepozitim " + str(int(round(float(data[6])))) + " meter katror. "
	return koment2

def get_rast(Kodi):
	import os
	csv_folder = r"F:\Dropbox\Projekte UTS-01\DLDP - Venddepozitime\Code"
	csv_file = r"Relacion.csv"
	my_data = readlines(os.path.join(csv_folder,csv_file))
	splitted_data = []
	for data in my_data:
		splitted_data.append(data.split(","))
	koment2 = ""
	for data in splitted_data:
		if str(data[0])== Kodi :
			if str(data[7]) == "1":
				koment2 = "Mbetjet ne kete venddepozitmim jane te rralla dhe rastesore. "
	return koment2

def get_bimesi(Kodi):
	import os
	csv_folder = r"F:\Dropbox\Projekte UTS-01\DLDP - Venddepozitime\Code"
	csv_file = r"Relacion.csv"
	my_data = readlines(os.path.join(csv_folder,csv_file))
	splitted_data = []
	for data in my_data:
		splitted_data.append(data.split(","))
	koment2 = ""
	for data in splitted_data:
		if str(data[0])==Kodi:
			if str(data[8]) == "1":
				koment2 = "Ky venddepozitim ndodhet ne nje siperfaqe territori ku gjenden peme perreth. Duhet te shikohet heqja e tyre nese do te vazhdoje te perdoret ky venddepozitim. "
			else:
				koment2 = "Perreth ketij venddepozitimi bimesia nuk paraqqitet shume e zhvilluar. "
	return koment2


def create_rel(kodi,folder,emri,x,y,bashkia,qarku):
	from docx import Document
	from docx.shared import Inches
	from docx.shared import Pt
	from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
	import os
	filename_word = os.path.join(folder,kodi) + ".docx"
	new_filename_word = os.path.join(folder,kodi) + "_.docx"
	table_point1 = "Komente për rezultatin e vend-depozitimit së mbetjeve. (Gjendja e vend-depozitimit / Risqet kryesore mjedisore / Risqet kryesore lidhur me popullatën)"
	table_point2 = "A ka ndonjë vend-hedhje mbeturinash afër? A ka ndonjë vend turistik të dukshëm nga vend-depozitimi e mbeturinave? Ose ndonjë zone të përcaktuar me VKM?"
	table_point3 = "Harta / Fotografia / Skicat e vend-depozitimit së mbeturinave"
	table_p0_text = [table_point1,table_point2,table_point3]
	doc_title = "Vend-depozitimi: %s" % (get_name(kodi),) # heading 1
	doc_normal_heading = "Rezultati final për Vend-depozitimin %s, i lokalizuar në koordinatat E: %s N: %s në Bashkinë %s, Qarku %s."
	
	text1 = ""
	text2 = ""
	Photos = []


	# doc_normal_heading style "Body text"
	
	# document.add_heading('Heading 1', level=1)
	get_table_from_xls_to_word(kodi, folder)
	document = Document(filename_word)
	h1_style = document.styles['Heading 1'].paragraph_format
	h1_style.space_after = Pt(25)
	bodyText_style = document.styles['Body Text Indent'].paragraph_format
	bodyText_style.space_after = Pt(15)
	h1 = document.add_paragraph(doc_title, style='Heading 1')
	p1 = document.add_paragraph(doc_normal_heading % (get_name(kodi), x, y, bashkia, qarku), style='Body Text Indent')
	# p1 = document.add_paragraph("Dummy text", style='Body Text Indent')
	p2 = document.add_paragraph("Tabela 1:	Rezultati i pikeve per venddepozitimin %s" % (emri,), style='Body Text Indent')
	table0 = document.tables[0]
	move_table_after(table0,p1)

	koment1 = ""
	koment1 = koment1 + get_sip(kodi)
	koment1 = koment1 + get_rast(kodi)
	koment1 = koment1 + get_bimesi(kodi)

	p2 = document.add_paragraph(table_point1 + "\n", style='Body Text Indent')
	p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
	p3 = document.add_paragraph(koment1 + "\n", style='Normal')
	p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
	p4 = document.add_paragraph(table_point2 + "\n", style='Body Text Indent')
	p4.alignment = WD_ALIGN_PARAGRAPH.LEFT

	csv_folder = r"F:\Dropbox\Projekte UTS-01\DLDP - Venddepozitime\Code"
	csv_file = r"Relacion.csv"
	my_data = readlines(os.path.join(csv_folder,csv_file))
	splitted_data = []
	for data in my_data:
		splitted_data.append(data.split(","))
	koment2 = ""
	for data in splitted_data:
		if str(data[0])==kodi:
			koment2 = data[1]
	if koment2 == None:
		koment2 = "Venddepozitimi nuk gjendet afer ndonje monumenti Kulture apo ne zone te mbrojtur."

	p5 = document.add_paragraph(koment2, style='Normal')
	# table1 = document.add_table(rows=3, cols=1)
	# # table1.style = 'LightShading-Accent1'
	# first_table_cell = table1.rows[0].cells[0]
	# p0 = first_table_cell.paragraphs[0]
	# p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
	# run = p0.add_run(table_point1 + "\n\n\n" + text1)
	# #run = p0.add_run(WD_BREAK.LINE)
	# #run = p0.add_run(table_point1)
	# first_table_cell = table1.rows[1].cells[0]
	# p0 = first_table_cell.paragraphs[0]
	# p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
	# run = p0.add_run(table_point2 + "\n\n\n" + text2)
	# first_table_cell = table1.rows[2].cells[0]
	# p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
	# p0 = first_table_cell.paragraphs[0]
	# run = p0.add_run(table_point3 + "\n\n\n")
	# pic_folder1 = r"F:\Dropbox\Projekte UTS-01\DLDP - Venddepozitime\WM-20180411T124506Z-001\Ministria e Turizmit dhe Mjedisit-WM\foto\Str.View"
	# pic_folder2 = r"F:\Dropbox\Projekte UTS-01\DLDP - Venddepozitime\WM-20180411T124506Z-001\Ministria e Turizmit dhe Mjedisit-WM\foto\Fotot VD"
	# pic_folder3 = r"F:\Dropbox\Projekte UTS-01\DLDP - Venddepozitime\WM-20180411T124506Z-001\Ministria e Turizmit dhe Mjedisit-WM\foto\Asig"
	# pic_folder4 = r"F:\Dropbox\Projekte UTS-01\DLDP - Venddepozitime\WM-20180411T124506Z-001\Ministria e Turizmit dhe Mjedisit-WM\foto\foto_id"
	# pic_folder_paths = [pic_folder1,pic_folder2,pic_folder3,pic_folder4]
	# extensions = [".jpeg", ".png", ".jpg"]
	# for extension in extensions:
	# 	pic_list = get_files_in_folder(pic_folder2,extension)
	# 	for image in pic_list:
	# 		if os.path.basename(image)[:len(kodi)] == kodi:
	# 			run.add_picture(image)
	# 			run.add_break(WD_BREAK.PAGE)
	# 			# p0.add_run('imazh nga google viti: %s' % (os.path.basename(image)[len(kodi):-len(extension)]))
	# 			run.add_break(WD_BREAK.PAGE)

	# 	pic_list = get_files_in_folder(pic_folder1,extension)
	# 	for image in pic_list:
	# 		if os.path.basename(image)[:-len(extension)+2] == kodi:
	# 			run.add_picture(image, width = Inches(1))
	# 			run.add_break(WD_BREAK.PAGE)
	# 			# p0.run.add_run('imazh nga google street view')
	# 			run.add_break(WD_BREAK.PAGE)
	# 			pic_list = get_files_in_folder(pic_folder1,extension)

	# 	pic_list = get_files_in_folder(pic_folder3,extension)
	# 	for image in pic_list:
	# 		if os.path.basename(image)[:-len(extension)+2] == kodi:
	# 			run.add_picture(image, width = Inches(1))
	# 			run.add_break(WD_BREAK.PAGE)
	# 			# p0.add_run('imazh nga MZHU')
	# 			run.add_break(WD_BREAK.PAGE)

	# 	pic_list = get_files_in_folder(pic_folder4,extension)
	# 	for image in pic_list:
	# 		if os.path.basename(image)[:-len(extension)+2] == kodi:
	# 			run.add_picture(image, width = Inches(1))
	# for extension in extensions:
	# 	for pic_folder in pic_folder_paths:
	# 		pic_list = get_files_in_folder(pic_folder,extension)
	# 		for image in pic_list:
	# 			# print(os.path.basename(image))
	# 			if os.path.basename(image).startswith(kodi):
	# 				# print("Entered run")
	# 				run.add_picture(image)
	document.save(new_filename_word)

def add_photos_to_rel(kodi,folder):
	from docx import Document
	from docx.shared import Inches
	from docx.shared import Pt
	from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
	import os
	new_filename_word = os.path.join(folder,kodi) + "_.docx"
	document = Document(new_filename_word)
	pic_folder1 = r"F:\Dropbox\Projekte UTS-01\DLDP - Venddepozitime\WM-20180411T124506Z-001\Ministria e Turizmit dhe Mjedisit-WM\foto\Str.View"
	pic_folder2 = r"F:\Dropbox\Projekte UTS-01\DLDP - Venddepozitime\WM-20180411T124506Z-001\Ministria e Turizmit dhe Mjedisit-WM\foto\Fotot VD"
	pic_folder3 = r"F:\Dropbox\Projekte UTS-01\DLDP - Venddepozitime\WM-20180411T124506Z-001\Ministria e Turizmit dhe Mjedisit-WM\foto\Asig"
	pic_folder4 = r"F:\Dropbox\Projekte UTS-01\DLDP - Venddepozitime\WM-20180411T124506Z-001\Ministria e Turizmit dhe Mjedisit-WM\foto\foto_id"
	pic_folder_paths = [pic_folder1,pic_folder2,pic_folder3,pic_folder4]
	extensions = [".jpeg", ".png", ".jpg"]
	for extension in extensions:
		pic_list = get_files_in_folder(pic_folder2,extension)
		for image in pic_list:
			if os.path.basename(image)[:-1*(len(extension)+2)] == kodi:
				# print(os.path.basename(image))
				document.add_picture(image, width=Inches(2.0))
				# p0.add_run('imazh nga google viti: %s' % (os.path.basename(image)[len(kodi):-len(extension)]))
				p4 = document.add_paragraph("Imazh nga google viti: 20%s" % (os.path.basename(image)[len(kodi):-1*(len(extension))],), style='Body Text Indent')
				p4.alignment = WD_ALIGN_PARAGRAPH.LEFT
	savefilename = os.path.join(folder,kodi) + "_imagesVD.docx"
	document.save(savefilename)

def get_google_image(kodi):
	import os
	pic_folder1 = r"F:\Dropbox\Projekte UTS-01\DLDP - Venddepozitime\WM-20180411T124506Z-001\Ministria e Turizmit dhe Mjedisit-WM\foto\Str.View"
	pic_folder2 = r"F:\Dropbox\Projekte UTS-01\DLDP - Venddepozitime\WM-20180411T124506Z-001\Ministria e Turizmit dhe Mjedisit-WM\foto\Fotot VD"
	pic_folder3 = r"F:\Dropbox\Projekte UTS-01\DLDP - Venddepozitime\WM-20180411T124506Z-001\Ministria e Turizmit dhe Mjedisit-WM\foto\Asig"
	pic_folder4 = r"F:\Dropbox\Projekte UTS-01\DLDP - Venddepozitime\WM-20180411T124506Z-001\Ministria e Turizmit dhe Mjedisit-WM\foto\foto_id"
	pic_folder_paths = [pic_folder1,pic_folder2,pic_folder3,pic_folder4]
	extensions = [".jpeg", ".png", ".jpg"]
	for extension in extensions:
		pic_list = get_files_in_folder(pic_folder2,extension)
		for image in pic_list:
			if os.path.basename(image)[:-1*(len(extension)+2)] == kodi:
				print(os.path.basename(image)[:-1*(len(extension)+2)])
				print(os.path.basename(image)[len(kodi):-1*(len(extension))])

def add_google_photos(kodi,folder):
	from docx import Document
	from docx.shared import Inches
	from docx.shared import Pt
	from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
	import os
	new_filename_word = os.path.join(folder,kodi) + "_.docx"
	document = Document(new_filename_word)
	pic_folder1 = r"F:\Dropbox\Projekte UTS-01\DLDP - Venddepozitime\WM-20180411T124506Z-001\Ministria e Turizmit dhe Mjedisit-WM\foto\Str.View"
	pic_folder2 = r"F:\Dropbox\Projekte UTS-01\DLDP - Venddepozitime\WM-20180411T124506Z-001\Ministria e Turizmit dhe Mjedisit-WM\foto\Fotot VD"
	pic_folder3 = r"F:\Dropbox\Projekte UTS-01\DLDP - Venddepozitime\WM-20180411T124506Z-001\Ministria e Turizmit dhe Mjedisit-WM\foto\Asig"
	pic_folder4 = r"F:\Dropbox\Projekte UTS-01\DLDP - Venddepozitime\WM-20180411T124506Z-001\Ministria e Turizmit dhe Mjedisit-WM\foto\foto_id"
	pic_folder_paths = [pic_folder1,pic_folder2,pic_folder3,pic_folder4]
	extensions = [".jpeg", ".png", ".jpg"]
	for extension in extensions:
		pic_list = get_files_in_folder(pic_folder2,extension)
		for image in pic_list:
			if os.path.basename(image)[:-1*(len(extension)+2)] == kodi:
				# print(os.path.basename(image))
				document.add_picture(image, width=Inches(5.0))
				# p0.add_run('imazh nga google viti: %s' % (os.path.basename(image)[len(kodi):-len(extension)]))
				p4 = document.add_paragraph("Imazh nga google viti: 20%s" % (os.path.basename(image)[len(kodi):-1*(len(extension))],), style='Body Text Indent')
				p4.alignment = WD_ALIGN_PARAGRAPH.LEFT
	savefilename = os.path.join(folder,kodi) + "_imagesVDG.docx"
	document.save(savefilename)

def add_streetview_photos(kodi,folder):
	from docx import Document
	from docx.shared import Inches
	from docx.shared import Pt
	from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
	import os
	new_filename_word = os.path.join(folder,kodi) + "_imagesVDG.docx"
	document = Document(new_filename_word)
	pic_folder1 = r"F:\Dropbox\Projekte UTS-01\DLDP - Venddepozitime\WM-20180411T124506Z-001\Ministria e Turizmit dhe Mjedisit-WM\foto\Str.View"
	pic_folder2 = r"F:\Dropbox\Projekte UTS-01\DLDP - Venddepozitime\WM-20180411T124506Z-001\Ministria e Turizmit dhe Mjedisit-WM\foto\Fotot VD"
	pic_folder3 = r"F:\Dropbox\Projekte UTS-01\DLDP - Venddepozitime\WM-20180411T124506Z-001\Ministria e Turizmit dhe Mjedisit-WM\foto\Asig"
	pic_folder4 = r"F:\Dropbox\Projekte UTS-01\DLDP - Venddepozitime\WM-20180411T124506Z-001\Ministria e Turizmit dhe Mjedisit-WM\foto\foto_id"
	pic_folder_paths = [pic_folder1,pic_folder2,pic_folder3,pic_folder4]
	extensions = [".jpeg", ".png", ".jpg"]
	for extension in extensions:
		pic_list = get_files_in_folder(pic_folder1,extension)
		for image in pic_list:
			if os.path.basename(image)[:-1*(len(extension)+2)] == kodi:
				# print(os.path.basename(image))
				document.add_picture(image, width=Inches(5.0))
				# p0.add_run('imazh nga google viti: %s' % (os.path.basename(image)[len(kodi):-len(extension)]))
				p4 = document.add_paragraph("Imazh nga google street view", style='Body Text Indent')
				p4.alignment = WD_ALIGN_PARAGRAPH.LEFT
	savefilename = os.path.join(folder,kodi) + "_imagesSTRV.docx"
	document.save(savefilename)
def add_asig_photos(kodi,folder):
	from docx import Document
	from docx.shared import Inches
	from docx.shared import Pt
	from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
	import os
	new_filename_word = os.path.join(folder,kodi) + "_imagesSTRV.docx"
	document = Document(new_filename_word)
	pic_folder1 = r"F:\Dropbox\Projekte UTS-01\DLDP - Venddepozitime\WM-20180411T124506Z-001\Ministria e Turizmit dhe Mjedisit-WM\foto\Str.View"
	pic_folder2 = r"F:\Dropbox\Projekte UTS-01\DLDP - Venddepozitime\WM-20180411T124506Z-001\Ministria e Turizmit dhe Mjedisit-WM\foto\Fotot VD"
	pic_folder3 = r"F:\Dropbox\Projekte UTS-01\DLDP - Venddepozitime\WM-20180411T124506Z-001\Ministria e Turizmit dhe Mjedisit-WM\foto\Asig"
	pic_folder4 = r"F:\Dropbox\Projekte UTS-01\DLDP - Venddepozitime\WM-20180411T124506Z-001\Ministria e Turizmit dhe Mjedisit-WM\foto\foto_id"
	pic_folder_paths = [pic_folder1,pic_folder2,pic_folder3,pic_folder4]
	extensions = [".jpeg", ".png", ".jpg"]
	for extension in extensions:
		pic_list = get_files_in_folder(pic_folder3,extension)
		for image in pic_list:
			if os.path.basename(image)[:-1*(len(extension)+2)] == kodi:
				# print(os.path.basename(image))
				document.add_picture(image, width=Inches(5.0))
				# p0.add_run('imazh nga google viti: %s' % (os.path.basename(image)[len(kodi):-len(extension)]))
				p4 = document.add_paragraph("Imazh nga website i Ministrisë së Zhvillimit Urban", style='Body Text Indent')
				p4.alignment = WD_ALIGN_PARAGRAPH.LEFT
	savefilename = os.path.join(folder,kodi) + "_imagesAsig.docx"
	document.save(savefilename)

def add_other_photos(kodi,folder):
	from docx import Document
	from docx.shared import Inches
	from docx.shared import Pt
	from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
	import os
	new_filename_word = os.path.join(folder,kodi) + "_.docx"
	document = Document(new_filename_word)
	pic_folder1 = r"F:\Dropbox\Projekte UTS-01\DLDP - Venddepozitime\WM-20180411T124506Z-001\Ministria e Turizmit dhe Mjedisit-WM\foto\Str.View"
	pic_folder2 = r"F:\Dropbox\Projekte UTS-01\DLDP - Venddepozitime\WM-20180411T124506Z-001\Ministria e Turizmit dhe Mjedisit-WM\foto\Fotot VD"
	pic_folder3 = r"F:\Dropbox\Projekte UTS-01\DLDP - Venddepozitime\WM-20180411T124506Z-001\Ministria e Turizmit dhe Mjedisit-WM\foto\Asig"
	pic_folder4 = r"F:\Dropbox\Projekte UTS-01\DLDP - Venddepozitime\WM-20180411T124506Z-001\Ministria e Turizmit dhe Mjedisit-WM\foto\foto_id"
	pic_folder_paths = [pic_folder1,pic_folder2,pic_folder3,pic_folder4]
	extensions = [".jpeg", ".png", ".jpg"]
	for extension in extensions:
		pic_list = get_files_in_folder(pic_folder4,extension)
		for image in pic_list:
			if os.path.basename(image).split("_")[0] == kodi:
				# print(os.path.basename(image))
				document.add_picture(image, width=Inches(5.0))
				# p0.add_run('imazh nga google viti: %s' % (os.path.basename(image)[len(kodi):-len(extension)]))
	savefilename = os.path.join(folder,kodi) + "_imagesTjera.docx"
	document.save(savefilename)


def getwordstyles(kodi,folder):
	import os
	from docx import Document
	filename_word = os.path.join(folder,kodi) + ".docx"
	document = Document(filename_word)
	print(document.styles)

'''
from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='IntenseQuote')

document.add_paragraph(
    'first item in unordered list', style='ListBullet'
)
document.add_paragraph(
    'first item in ordered list', style='ListNumber'
)

document.add_picture('monty-truth.png', width=Inches(1.25))

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for item in recordset:
    row_cells = table.add_row().cells
    row_cells[0].text = str(item.qty)
    row_cells[1].text = str(item.id)
    row_cells[2].text = item.desc

document.add_page_break()

document.save('demo.docx')
'''
