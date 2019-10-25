'''
Example on adding photo on word tab using docx

import docx
pic = "some_picture.jpg"
document = docx.Document()
tbl = document.add_table(rows=1, cols=1)
row_cells = tbl.add_row().cells
paragraph = row_cells[0].paragraphs[0]
run = paragraph.add_run()
run.add_picture(pic, width = 1400000, height = 1400000)
document.save("demo.docx")
'''
